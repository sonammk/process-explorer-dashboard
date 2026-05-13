from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Process Explorer Tool Demonstration Report.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = (Inches(2.0), Inches(4.25))
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = widths[0]
        cells[1].width = widths[1]
        set_cell_text(cells[0], label, bold=True, color="1F4E79")
        set_cell_text(cells[1], value)
    return table


def add_rubric(doc):
    doc.add_heading("Rubric Mapping", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["S.No", "Evaluation Criteria", "How This Report Covers It", "Max Marks"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[idx], "1F4E79")

    rows = [
        ("1", "Introduction of the Tool", "Tool name, type, category, platform, and purpose are clearly stated.", "2"),
        ("2", "Step-by-Step Demonstration", "The demonstration workflow explains how to launch, inspect, search, and analyze processes.", "3"),
        ("3", "Interpretation of Output", "CPU, memory, threads, handles, DLLs, process tree, and colors are interpreted.", "2"),
        ("4", "Team Coordination & Presentation", "Live demo sections can be divided among presenters using the steps and observations.", "2"),
        ("5", "Report", "Structured report follows the given activity format.", "1"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(68, 68, 68)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Report for the Tool Demonstration")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Process Explorer (Sysinternals)")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Category: Process Management | Type: GUI Tool | Platform: Windows")

    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Tool Name", "Process Explorer"),
            ("Developer", "Microsoft Sysinternals"),
            ("OS Concept", "Process management, CPU scheduling, memory usage, handles, threads, DLLs, and parent-child process relationships"),
            ("Best Demo Use", "Showing how the operating system creates, schedules, tracks, and terminates running processes"),
        ],
    )

    doc.add_heading("Index Sheet", level=1)
    for item in [
        "Introduction to the Tool",
        "Purpose of the Tool",
        "Features of the Tool",
        "Working Principle / OS Concept Behind It",
        "Steps to Use the Tool",
        "Observations and Output Interpretation",
        "Conclusion",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    doc.add_heading("1. Introduction to the Tool", level=1)
    doc.add_paragraph(
        "Process Explorer is an advanced graphical process-management utility from Microsoft Sysinternals. "
        "It provides a detailed live view of running processes, their parent-child relationships, CPU usage, "
        "memory consumption, threads, loaded DLLs, open handles, and security-related information."
    )
    add_bullet(doc, "Name of the tool: Process Explorer")
    add_bullet(doc, "Category: Process management and operating-system monitoring")
    add_bullet(doc, "Type: Graphical user interface utility")
    add_bullet(doc, "Platform: Windows")

    doc.add_heading("2. Purpose of the Tool", level=1)
    doc.add_paragraph(
        "The tool is used to observe and analyze processes in real time. It helps users understand what programs "
        "are running, how much CPU and memory each process consumes, which files or registry keys a process has "
        "opened, and how processes are related to each other."
    )
    for item in [
        "Demonstrates process creation, execution, suspension, and termination.",
        "Shows CPU scheduling effects through changing CPU usage values.",
        "Connects memory-management concepts to private bytes, working set, and virtual memory statistics.",
        "Reveals inter-process relationships through a process tree.",
        "Helps identify resource-heavy, suspicious, or unresponsive processes.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. Features of the Tool", level=1)
    for item in [
        "Tree view showing parent and child processes.",
        "Real-time CPU, memory, I/O, and GPU-related process information.",
        "Color-coded process categories for quick visual identification.",
        "Lower pane that displays handles or loaded DLLs for the selected process.",
        "Process properties window with tabs for image path, performance, threads, TCP/IP, security, environment, and strings.",
        "Search option for finding which process has opened a file, folder, DLL, or handle.",
        "Options to kill, suspend, restart, or set priority for a process.",
        "VirusTotal integration for checking suspicious executables, when enabled by the user.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4. Working Principle / OS Concept Behind It", level=1)
    doc.add_paragraph(
        "Process Explorer reads information exposed by Windows kernel and system APIs. The operating system keeps "
        "tables and structures for every process and thread. Process Explorer queries those structures and presents "
        "them in a readable dashboard."
    )
    for item in [
        "Process table: shows process IDs, names, users, and parent-child relationships.",
        "Scheduler data: shows CPU percentage and thread activity.",
        "Memory manager data: shows working set, private bytes, and virtual memory usage.",
        "Object manager data: shows handles for files, registry keys, events, and synchronization objects.",
        "Loader information: shows DLLs and modules loaded inside a process.",
        "Security subsystem: shows user account, integrity level, permissions, and digital signature details.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. Steps to Use the Tool", level=1)
    for step in [
        "Download Process Explorer from the Microsoft Sysinternals website.",
        "Extract the downloaded ZIP file and run procexp.exe or procexp64.exe.",
        "Accept the license agreement when launching it for the first time.",
        "Observe the main process tree and identify system processes, user applications, and background services.",
        "Click a process such as a browser, file explorer, or text editor and review CPU, memory, PID, and company name.",
        "Open View > Lower Pane View and choose Handles or DLLs to inspect resources used by the selected process.",
        "Double-click a process to open Properties and demonstrate the Performance, Threads, TCP/IP, and Security tabs.",
        "Use Find > Find Handle or DLL to search for an open file or loaded module.",
        "Right-click a non-critical test application and demonstrate safe actions such as Properties or Set Priority. Avoid terminating system processes during a live demo.",
        "Summarize how the displayed metrics connect to process management, scheduling, memory management, and resource allocation.",
    ]:
        add_numbered(doc, step)

    doc.add_heading("6. Observations and Output Interpretation", level=1)
    observations = [
        ("Process Tree", "Shows how applications are started by parent processes. Child processes are indented below their parent."),
        ("PID", "A unique process identifier assigned by the operating system to track each running process."),
        ("CPU %", "Shows how much processor time the process is currently using. High values indicate active computation."),
        ("Private Bytes", "Memory allocated exclusively for that process and not shared with other processes."),
        ("Working Set", "Physical memory currently used by the process."),
        ("Threads", "Execution units inside a process. More threads can mean parallel or background work."),
        ("Handles", "References to OS-managed objects such as files, registry keys, events, and sockets."),
        ("DLLs", "Shared libraries loaded into the process address space."),
        ("Colors", "Color highlighting helps distinguish services, packed images, jobs, new processes, and terminated processes."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Output / Metric", bold=True, color="FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Interpretation", bold=True, color="FFFFFF")
    set_cell_shading(table.rows[0].cells[0], "1F4E79")
    set_cell_shading(table.rows[0].cells[1], "1F4E79")
    for metric, interpretation in observations:
        cells = table.add_row().cells
        set_cell_text(cells[0], metric, bold=True)
        set_cell_text(cells[1], interpretation)

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Process Explorer is a powerful tool for demonstrating process management in Windows. It shows how the "
        "operating system organizes processes, schedules CPU time, assigns memory, manages handles, loads DLLs, "
        "and enforces security information. Compared with the normal Task Manager, it gives deeper visibility into "
        "what is happening behind the scenes, making it useful for both learning OS concepts and troubleshooting."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_rubric(doc)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
